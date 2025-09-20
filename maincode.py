# E-commerce Furniture Sales Prediction (Final Version with Saved Plots + Word Report)

import os
import re
import math
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt

from sklearn.model_selection import train_test_split, KFold, cross_val_score
from sklearn.preprocessing import OneHotEncoder, StandardScaler
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.compose import ColumnTransformer
from sklearn.linear_model import LinearRegression, Ridge
from sklearn.ensemble import RandomForestRegressor
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
from scipy.sparse import hstack
import joblib
from docx import Document
from docx.shared import Inches

# 1. Setup folders
os.makedirs("outputs/plots", exist_ok=True)
os.makedirs("outputs", exist_ok=True)

# 2. Load dataset
DATA_PATH = "ecommerce_furniture_dataset_2024.csv"
df = pd.read_csv(DATA_PATH)
print("Raw shape:", df.shape)

# 3. Cleaning & feature engineering
def coerce_price(x):
    if pd.isna(x):
        return np.nan
    s = str(x)
    s = re.sub(r"\s+", "", s)
    s = re.sub(r"[^0-9.]", "", s)
    if s.count(".") > 1:
        parts = s.split(".")
        s = parts[0] + "." + "".join(parts[1:])
    try:
        return float(s) if s else np.nan
    except:
        return np.nan

df["price_clean"] = df["price"].apply(coerce_price)
df["originalPrice_clean"] = df["originalPrice"].apply(coerce_price)
df["sold_num"] = pd.to_numeric(df["sold"], errors="coerce")

tag = df["tagText"].astype(str).str.strip()
tag = tag.replace({"Free  shipping": "Free shipping", "free shipping": "Free shipping"})
top_tags = tag.value_counts().head(5).index
df["tagText_simple"] = np.where(tag.isin(top_tags), tag, "Other")

df["discount_pct"] = (df["originalPrice_clean"] - df["price_clean"]) / df["originalPrice_clean"] * 100
df = df.dropna(subset=["price_clean", "sold_num"])
df["discount_pct"] = df["discount_pct"].fillna(0)

print("After cleaning:", df.shape)

# 4. Save EDA plots
plot_files = []

plt.hist(df["price_clean"], bins=50)
plt.title("Distribution of Price")
plt.xlabel("Price")
plt.ylabel("Count")
plot_path = "outputs/plots/price_distribution.png"
plt.savefig(plot_path)
plt.close()
plot_files.append(plot_path)

plt.hist(df["sold_num"], bins=50)
plt.title("Distribution of Sold")
plt.xlabel("Units Sold")
plt.ylabel("Count")
plot_path = "outputs/plots/sold_distribution.png"
plt.savefig(plot_path)
plt.close()
plot_files.append(plot_path)

plt.scatter(df["price_clean"], df["sold_num"], alpha=0.5)
plt.title("Price vs Sold")
plt.xlabel("Price")
plt.ylabel("Sold")
plot_path = "outputs/plots/price_vs_sold.png"
plt.savefig(plot_path)
plt.close()
plot_files.append(plot_path)

# 5. Features & Train/Test split
feature_cols_num = ["price_clean", "discount_pct"]
feature_cols_cat = ["tagText_simple"]
text_col = "productTitle"

X = df[feature_cols_num + feature_cols_cat + [text_col]]
y = df["sold_num"]

preproc_tab = ColumnTransformer(
    transformers=[
        ("num", StandardScaler(), feature_cols_num),
        ("cat", OneHotEncoder(handle_unknown="ignore"), feature_cols_cat),
    ],
    remainder="drop"
)

X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

Xt_train_tab = preproc_tab.fit_transform(X_train)
Xt_test_tab = preproc_tab.transform(X_test)

tfidf = TfidfVectorizer(max_features=300, ngram_range=(1,2))
Xt_train_txt = tfidf.fit_transform(X_train[text_col].astype(str))
Xt_test_txt = tfidf.transform(X_test[text_col].astype(str))

Xt_train = hstack([Xt_train_tab, Xt_train_txt])
Xt_test = hstack([Xt_test_tab, Xt_test_txt])

# 6. Train Models
models = {
    "LinearRegression": LinearRegression(),
    "Ridge": Ridge(alpha=1.0),
    "RandomForest": RandomForestRegressor(n_estimators=300, random_state=42, n_jobs=-1)
}

results = {}
for name, model in models.items():
    model.fit(Xt_train, y_train)
    y_pred = model.predict(Xt_test)
    mse = mean_squared_error(y_test, y_pred)
    rmse = math.sqrt(mse)
    mae = mean_absolute_error(y_test, y_pred)
    r2 = r2_score(y_test, y_pred)
    results[name] = {"RMSE": rmse, "MAE": mae, "R2": r2}
    print(f"{name:>16} -> RMSE: {rmse:.2f} | MAE: {mae:.2f} | R2: {r2:.4f}")

best_name = min(results, key=lambda k: results[k]["RMSE"])
best_model = models[best_name]
print("\nBest model:", best_name)

# 7. Cross-validation
kf = KFold(n_splits=5, shuffle=True, random_state=42)
cv_scores = cross_val_score(best_model, Xt_train, y_train, cv=kf, scoring="r2", n_jobs=-1)
print("CV R2 scores:", cv_scores)
print("CV R2 mean:", cv_scores.mean())

# 8. Save best model
bundle = {
    "model_name": best_name,
    "model": best_model,
    "preproc_tab": preproc_tab,
    "tfidf": tfidf
}
joblib.dump(bundle, "outputs/rf_sales_model.pkl")
print("Saved best model as outputs/rf_sales_model.pkl")

# 9. Generate Word Report with plots
OUTPUT_DOCX = "outputs/furniture_sales_report.docx"
n_rows, n_cols = df.shape
missing_vals = df.isnull().sum().to_dict()

doc = Document()
doc.add_heading("E-commerce Furniture Sales Prediction Report", 0)

doc.add_heading("Dataset Overview", level=1)
doc.add_paragraph(f"Total Rows: {n_rows}")
doc.add_paragraph(f"Total Columns: {n_cols}")

doc.add_heading("Missing Values per Column", level=2)
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text, hdr_cells[1].text = "Column", "Missing Count"
for k, v in missing_vals.items():
    row_cells = table.add_row().cells
    row_cells[0].text = str(k)
    row_cells[1].text = str(v)

doc.add_heading("Exploratory Data Analysis (EDA)", level=1)
for plot in plot_files:
    doc.add_paragraph(os.path.basename(plot).replace(".png", "").replace("_", " ").title())
    doc.add_picture(plot, width=Inches(5.5))
    doc.add_paragraph("")

doc.add_heading("Model Evaluation Results", level=1)
table2 = doc.add_table(rows=1, cols=4)
hdr2 = table2.rows[0].cells
hdr2[0].text, hdr2[1].text, hdr2[2].text, hdr2[3].text = "Model", "RMSE", "MAE", "R²"
for model_name, metrics in results.items():
    row_cells = table2.add_row().cells
    row_cells[0].text = model_name
    row_cells[1].text = f"{metrics['RMSE']:.2f}"
    row_cells[2].text = f"{metrics['MAE']:.2f}"
    row_cells[3].text = f"{metrics['R2']:.3f}" + (" (Best)" if model_name == best_name else "")

doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    f"The {best_name} performed the best on the dataset.\n"
    "It can be used to predict the number of furniture items sold based on product attributes "
    "such as price, discount percentage, product title, and shipping tags.\n\n"
    "This project demonstrates a full ML pipeline: data cleaning, exploratory analysis, "
    "feature engineering, model training, evaluation, and deployment."
)

doc.save(OUTPUT_DOCX)
print(f"✅ Report saved as {OUTPUT_DOCX}")

# 10. Prediction helper
def predict_sold(bundle_path, new_data: pd.DataFrame):
    bundle = joblib.load(bundle_path)
    model = bundle["model"]
    preproc_tab = bundle["preproc_tab"]
    tfidf = bundle["tfidf"]

    new_data["price_clean"] = new_data["price"].apply(coerce_price)
    new_data["originalPrice_clean"] = new_data["originalPrice"].apply(coerce_price)
    new_data["discount_pct"] = (new_data["originalPrice_clean"] - new_data["price_clean"]) / new_data["originalPrice_clean"] * 100
    new_data["discount_pct"] = new_data["discount_pct"].fillna(0)

    tag = new_data["tagText"].astype(str).str.strip()
    new_data["tagText_simple"] = np.where(tag.isin(["Free shipping"]), "Free shipping", "Other")

    Xt_tab = preproc_tab.transform(new_data[["price_clean", "discount_pct", "tagText_simple"]])
    Xt_txt = tfidf.transform(new_data["productTitle"].astype(str))
    X_final = hstack([Xt_tab, Xt_txt])

    return model.predict(X_final)

# Example predictions
new_products = pd.DataFrame([
    {"productTitle": "Modern Wooden Desk with Drawers", "price": "$149.99", "originalPrice": "$189.99", "tagText": "Free shipping"},
    {"productTitle": "Luxury Leather Sofa Set", "price": "$999.00", "originalPrice": "$1299.00", "tagText": "Free shipping"}
])
preds = predict_sold("outputs/rf_sales_model.pkl", new_products)
print("\nPredictions for new products:")
for title, pred in zip(new_products["productTitle"], preds):
    print(f"{title[:40]}... -> Predicted Sold: {int(pred)} units")